"""Multi-format document parser: PDF, DOCX, Markdown, HTML -> plain text.

Follows the project pattern of graceful degradation:
- Each format parser is optional (missing lib -> skip format).
- Returns structured ParsedDocument with text + metadata.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_TYPES = {"pdf", "docx", "md", "html", "txt"}


@dataclass
class ParsedPage:
    """A single page / section of a document."""
    page_num: int
    text: str


@dataclass
class ParsedDocument:
    """Result of parsing a document file."""
    filename: str
    file_type: str
    pages: list[ParsedPage]
    full_text: str = ""
    page_count: int = 0
    metadata: dict = field(default_factory=dict)
    errors: list[str] = field(default_factory=list)

    def __post_init__(self) -> None:
        if not self.full_text:
            self.full_text = "\n\n".join(p.text for p in self.pages)
        if not self.page_count:
            self.page_count = len(self.pages)


def detect_file_type(filename: str) -> str | None:
    """Detect file type from filename extension."""
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext in ("md", "markdown"):
        return "md"
    if ext in SUPPORTED_TYPES:
        return ext
    return None


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    """Parse a document file and extract text.

    Dispatches to format-specific parsers based on file extension.
    """
    file_type = detect_file_type(filename)
    if file_type is None:
        return ParsedDocument(
            filename=filename,
            file_type="unknown",
            pages=[],
            errors=[f"Unsupported file type: {Path(filename).suffix}"],
        )

    parser_map = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "md": _parse_markdown,
        "html": _parse_html,
        "txt": _parse_txt,
    }

    parser = parser_map.get(file_type)
    if parser is None:
        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            pages=[],
            errors=[f"No parser available for {file_type}"],
        )

    try:
        return parser(filename, content)
    except Exception as e:
        logger.error("Failed to parse %s: %s", filename, e, exc_info=True)
        return ParsedDocument(
            filename=filename,
            file_type=file_type,
            pages=[],
            errors=[f"Parse error: {e}"],
        )


def _parse_pdf(filename: str, content: bytes) -> ParsedDocument:
    """Parse PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        # Fallback to pdfplumber
        return _parse_pdf_pdfplumber(filename, content)

    doc = fitz.open(stream=content, filetype="pdf")
    pages: list[ParsedPage] = []
    metadata = {}

    # Extract metadata
    pdf_meta = doc.metadata or {}
    if pdf_meta.get("title"):
        metadata["title"] = pdf_meta["title"]
    if pdf_meta.get("author"):
        metadata["author"] = pdf_meta["author"]

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages.append(ParsedPage(page_num=page_num + 1, text=text.strip()))

    doc.close()
    return ParsedDocument(
        filename=filename,
        file_type="pdf",
        pages=pages,
        metadata=metadata,
    )


def _parse_pdf_pdfplumber(filename: str, content: bytes) -> ParsedDocument:
    """Fallback PDF parser using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            pages=[],
            errors=["No PDF library available (install PyMuPDF or pdfplumber)"],
        )

    import io
    pages: list[ParsedPage] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(ParsedPage(page_num=i + 1, text=text.strip()))

    return ParsedDocument(
        filename=filename,
        file_type="pdf",
        pages=pages,
    )


def _parse_docx(filename: str, content: bytes) -> ParsedDocument:
    """Parse DOCX using python-docx."""
    try:
        import docx
    except ImportError:
        return ParsedDocument(
            filename=filename,
            file_type="docx",
            pages=[],
            errors=["python-docx not installed"],
        )

    import io
    doc = docx.Document(io.BytesIO(content))
    metadata = {}

    # Extract metadata from core properties
    props = doc.core_properties
    if props.title:
        metadata["title"] = props.title
    if props.author:
        metadata["author"] = props.author

    # DOCX doesn't have pages -> treat entire document as one "page"
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    pages = [ParsedPage(page_num=1, text=full_text)] if full_text.strip() else []

    return ParsedDocument(
        filename=filename,
        file_type="docx",
        pages=pages,
        metadata=metadata,
    )


def _parse_markdown(filename: str, content: bytes) -> ParsedDocument:
    """Parse Markdown -> plain text (strip formatting)."""
    text = _decode_text(content)
    if text is None:
        return ParsedDocument(
            filename=filename,
            file_type="md",
            pages=[],
            errors=["Unable to decode markdown file"],
        )

    # Simple markdown stripping (no external dep required)
    plain = _strip_markdown(text)
    pages = [ParsedPage(page_num=1, text=plain)] if plain.strip() else []

    return ParsedDocument(
        filename=filename,
        file_type="md",
        pages=pages,
    )


def _parse_html(filename: str, content: bytes) -> ParsedDocument:
    """Parse HTML -> plain text."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # Fallback: regex-based tag stripping
        text = _decode_text(content) or ""
        plain = re.sub(r"<[^>]+>", " ", text)
        plain = re.sub(r"\s+", " ", plain).strip()
        pages = [ParsedPage(page_num=1, text=plain)] if plain else []
        return ParsedDocument(filename=filename, file_type="html", pages=pages)

    text = _decode_text(content) or ""
    soup = BeautifulSoup(text, "html.parser")

    # Remove script and style elements
    for tag in soup(["script", "style"]):
        tag.decompose()

    plain = soup.get_text(separator="\n")
    # Collapse multiple blank lines
    plain = re.sub(r"\n{3,}", "\n\n", plain).strip()
    pages = [ParsedPage(page_num=1, text=plain)] if plain else []

    return ParsedDocument(
        filename=filename,
        file_type="html",
        pages=pages,
    )


def _parse_txt(filename: str, content: bytes) -> ParsedDocument:
    """Parse plain text file."""
    text = _decode_text(content)
    if text is None:
        return ParsedDocument(
            filename=filename,
            file_type="txt",
            pages=[],
            errors=["Unable to decode text file"],
        )

    pages = [ParsedPage(page_num=1, text=text.strip())] if text.strip() else []
    return ParsedDocument(
        filename=filename,
        file_type="txt",
        pages=pages,
    )


def _decode_text(content: bytes) -> str | None:
    """Try decoding bytes with multiple encodings (UTF-8 first, then Korean)."""
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return None


def _strip_markdown(text: str) -> str:
    """Strip common markdown formatting to get plain text."""
    # Remove headers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove bold/italic
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}(.+?)_{1,3}", r"\1", text)
    # Remove links [text](url)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Remove images ![alt](url)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    # Remove code blocks
    text = re.sub(r"```[\s\S]*?```", "", text)
    # Remove inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Remove blockquotes
    text = re.sub(r"^>\s+", "", text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Clean up extra whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
